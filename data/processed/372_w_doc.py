import re
import glob
from docx import Document


def task_func(directory_path: str) -> int:
    """
    Processes all Word (.docx) files in the provided directory, searching for double quotes in the text 
    and adding a backslash before each double quote to "protect" it.
    
    Parameters:
    - directory_path (str): Path to the directory containing .docx files to be processed.
    
    Returns:
    - int: Number of .docx files processed.

    Requirements:
    - re
    - docx
    - glob

    Example:
    >>> import tempfile
    >>> temp_dir = tempfile.mkdtemp()
    >>> doc = Document()
    >>> _ = doc.add_paragraph("This is a sample text with double quotes.")
    >>> doc.save(temp_dir + '/sample.docx')
    >>> task_func(temp_dir)
    1
    """
    docx_files = glob.glob(directory_path + '/*.docx')
    processed_files = 0
    for docx_file in docx_files:
        document = Document(docx_file)
        for paragraph in document.paragraphs:
            paragraph.text = re.sub(r'(?<!\\)"', r'\"', paragraph.text)
        document.save(docx_file)
        processed_files += 1
    return processed_files

import unittest
import shutil
import os
import doctest
import tempfile
class TestCases(unittest.TestCase):
    def setUp(self):
        self.base_tmp_dir = tempfile.mkdtemp()
        self.test_directory = f"{self.base_tmp_dir}/test/"
        if not os.path.exists(self.test_directory):
            os.makedirs(self.test_directory)
        test_data = {
            "file_1.docx": "This is a sample text without any double quotes.",
            "file_2.docx": "This is a \"sample\" text with double quotes.",
            "file_3.docx": r'This is a \"sample\" text with double quotes already protected.',
            "file_4.docx": "Hello \"world\"! How are you \"today\"?",
            "file_5.docx": "Testing \"multiple\" paragraphs.\n\nAnother paragraph with \"quotes\"."
        }
        # Create .docx files for each scenario
        for file_name, content in test_data.items():
            doc = Document()
            for paragraph in content.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(self.test_directory + file_name)
    def tearDown(self):
        if os.path.exists(self.test_directory):
            shutil.rmtree(self.test_directory)
    def read_docx_content(self, file_path):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def test_case_1(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_1.docx")
        self.assertEqual(content, "This is a sample text without any double quotes.")
    def test_case_2(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_2.docx")
        self.assertEqual(content, r'This is a \"sample\" text with double quotes.')
    def test_case_3(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_3.docx")
        self.assertEqual(content, r'This is a \"sample\" text with double quotes already protected.')
    def test_case_4(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_4.docx")
        self.assertEqual(content, r'Hello \"world\"! How are you \"today\"?')
    def test_case_5(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_5.docx")
        self.assertEqual(content, 'Testing \\"multiple\\" paragraphs.\n\nAnother paragraph with \\"quotes\\".')
